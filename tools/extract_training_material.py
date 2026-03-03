from __future__ import annotations

import json
import html
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

try:
    import fitz  # type: ignore
    import cv2  # type: ignore
    import numpy as np  # type: ignore
    from rapidocr_onnxruntime import RapidOCR  # type: ignore
except Exception:  # pragma: no cover - optional OCR dependencies
    fitz = None
    cv2 = None
    np = None
    RapidOCR = None

REPO_ROOT = Path(__file__).resolve().parents[1]

SOURCE_FILES = [
    Path(r"c:\MSSA Training\MSSA_CAD_Ultimate_Guide_22_Jan_2026.pdf"),
    Path(r"c:\MSSA Training\az2006_portrait_dark_teal.pdf"),
    Path(r"c:\MSSA Training\MSSA_Comprehensive_Reference_no_strike.pdf"),
    Path(r"c:\MSSA Training\AI900_Cheatsheet.pdf"),
    Path(r"c:\MSSA Training\MSSA Training App.docx"),
    Path(r"c:\MSSA Training\Vibe_Check_Flow_Full_Instructions.pdf"),
    Path(r"c:\MSSA Training\C# code snippets - Visual Studio (Windows) _ Microsoft Learn.pdf"),
    Path(r"c:\MSSA Training\keyboard-shortcuts-windows.pdf"),
    Path(r"c:\MSSA Training\NumbersAndMemory.xlsx"),
    Path(r"c:\MSSA Training\Int, Var, Methodologies, & Errors - All-in-one.pdf"),
    Path(r"c:\MSSA Training\Comprehensive_Reference.pdf"),
    Path(r"c:\MSSA Training\Coding Commands Master Reference - Java Script - Python - C#.pdf"),
    Path(r"c:\MSSA Training\Layouts_and_Instances_Deep_Guide.pdf"),
    Path(r"c:\MSSA Training\Git_Markdown_Master_Reference.pdf"),
    Path(r"c:\MSSA Training\Training Notes\Daily Training Notes\Training Notes (11-26 Feb 2026).pdf"),
]

# Handle smart punctuation variants present in Windows paths.
ALT_SOURCE_FILES = [
    Path(r"c:\MSSA Training\Int, Var, Methodologies, & Errors - All‑in‑one.pdf"),
    Path(r"c:\MSSA Training\Int, Var, Methodologies, & Errors — All‑in‑one.pdf"),
    Path(r"c:\MSSA Training\Coding Commands Master Reference — Java Script • Python • C#.pdf"),
    Path(r"c:\MSSA Training\Coding Commands Master Reference - Java Script - Python - C#.pdf"),
]

OUTPUT_WEB_SOURCE = REPO_ROOT / "StudyTide Forge Web" / "seed-source" / "legacy-source.cs"
OUTPUT_DESKTOP_SOURCE = REPO_ROOT / "StudyTide Forge" / "seed-source" / "legacy-source.cs"
OUTPUT_JSON = REPO_ROOT / "App_Data" / "training-material" / "extracted-training-pairs.json"
OUTPUT_REPORT = REPO_ROOT / "App_Data" / "training-material" / "extraction-report.md"

NOISE_PATTERNS = [
    re.compile(r"^page\s+\d+(\s+of\s+\d+)?$", re.IGNORECASE),
    re.compile(r"^mssa\s+cad\s+ultimate\s+guide.*$", re.IGNORECASE),
    re.compile(r"^all\s+rights\s+reserved.*$", re.IGNORECASE),
    re.compile(r"^printed\s+or\s+handwritten\s+text.*$", re.IGNORECASE),
    re.compile(r"^table\s+of\s+contents$", re.IGNORECASE),
    re.compile(r"^introduction$", re.IGNORECASE),
    re.compile(r"^name\s+or\s+shortcut$", re.IGNORECASE),
    re.compile(r"^format:\s*command", re.IGNORECASE),
]

GENERIC_TERMS = {
    "introduction",
    "overview",
    "summary",
    "general",
    "notes",
    "definitions",
    "workflow",
    "example",
    "examples",
    "output",
    "input",
    "syntax",
    "tips",
    "note",
    "notes",
}


@dataclass(frozen=True)
class TrainingPair:
    question: str
    answer: str
    source_file: str


@dataclass
class SourceStats:
    source_file: str
    pages_or_sections: int
    extracted_lines: int
    qualified_statements: int
    generated_pairs: int
    used_ocr: bool


def resolve_source_files() -> list[Path]:
    files: list[Path] = []

    for source in SOURCE_FILES:
        if source.exists():
            files.append(source)

    for candidate in ALT_SOURCE_FILES:
        if candidate.exists() and candidate not in files:
            files.append(candidate)

    # Stable order by full path keeps generation deterministic.
    files = sorted(files, key=lambda p: str(p).lower())
    return files


def normalize_ascii(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    return normalized.encode("ascii", "ignore").decode("ascii")


def normalize_space(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def normalize_for_key(value: str) -> str:
    value = normalize_ascii(value).lower()
    value = re.sub(r"[^a-z0-9\s]+", " ", value)
    return normalize_space(value)


def clean_line(raw: str) -> str:
    text = html.unescape(raw)
    text = re.sub(r"<[^>]+>", " ", text)
    text = text.replace("\x00", " ")
    text = text.replace("\u25a0", " ")
    text = text.replace("\u2013", "-")
    text = text.replace("\u2014", "-")
    text = text.replace("\u2018", "'")
    text = text.replace("\u2019", "'")
    text = text.replace("\u201c", '"')
    text = text.replace("\u201d", '"')
    text = normalize_ascii(text)
    text = normalize_space(text)
    text = re.sub(r"^[\-\*\u2022\u00b7\+]+\s*", "", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def strip_reference_noise(value: str) -> str:
    value = re.sub(r"\b\d{10,}\b", " ", value)
    value = re.sub(r"\bL\d{1,4}-L\d{1,4}\b", " ", value)
    value = re.sub(r"\b\d{1,4}:\d{1,4}\b", " ", value)
    return normalize_space(value)


def is_noise(line: str) -> bool:
    if not line:
        return True

    lowered = line.lower().strip()

    for pattern in NOISE_PATTERNS:
        if pattern.match(lowered):
            return True

    if lowered.startswith("http://") or lowered.startswith("https://"):
        return True

    if len(lowered) < 8:
        return True

    if lowered.startswith("#") and len(lowered.split()) <= 4:
        return True

    alpha_count = sum(1 for char in lowered if char.isalpha())
    digit_count = sum(1 for char in lowered if char.isdigit())
    punctuation_count = sum(1 for char in lowered if not char.isalnum() and not char.isspace())

    if alpha_count < 3:
        return True

    if digit_count > 0 and alpha_count / max(1, len(lowered)) < 0.2:
        return True

    if punctuation_count / max(1, len(lowered)) > 0.25:
        return True

    return False


def is_heading(line: str) -> bool:
    if not line:
        return False

    if len(line) > 92:
        return False

    if re.search(r"[.!?;]$", line):
        return False

    words = re.findall(r"[A-Za-z0-9#+./()'-]+", line)
    if len(words) == 0 or len(words) > 10:
        return False

    lowered = line.lower()
    if lowered.startswith("when ") or lowered.startswith("use ") or lowered.startswith("avoid "):
        return False

    if ":" in line and len(line.split(":", 1)[1].strip()) > 16:
        return False

    if re.search(r"\b(ctrl|alt|shift)\+", lowered):
        return False

    if re.search(r"\b\d{4,}\b", lowered):
        return False

    if any(token in line for token in ("{", "}", ";", "//", "\"", "=")):
        return False

    return True


def should_join(previous: str, current: str) -> bool:
    if not previous or not current:
        return False

    if len(previous) + len(current) > 380:
        return False

    if previous.endswith(":") or previous.endswith(",") or previous.endswith("-"):
        return True

    if re.search(r"[.!?]$", previous):
        return False

    if current[0].islower() or current.startswith("("):
        return True

    if len(previous.split()) >= 10:
        return True

    return False


def coalesce_lines(lines: Iterable[str]) -> list[str]:
    merged: list[str] = []
    buffer = ""

    for line in lines:
        if not buffer:
            buffer = line
            continue

        if should_join(buffer, line):
            buffer = normalize_space(f"{buffer} {line}")
        else:
            merged.append(buffer)
            buffer = line

    if buffer:
        merged.append(buffer)

    return merged


def extract_pdf_lines(path: Path) -> tuple[list[str], int, bool]:
    reader = PdfReader(str(path))
    raw_lines: list[str] = []
    used_ocr = False

    blank_pages: list[int] = []

    for page_index, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        cleaned = [clean_line(line) for line in text.splitlines()]
        cleaned = [line for line in cleaned if line]

        if len("".join(cleaned)) < 24:
            blank_pages.append(page_index)
        else:
            raw_lines.extend(cleaned)

    if blank_pages and fitz is not None and RapidOCR is not None and cv2 is not None and np is not None:
        used_ocr = True
        doc = fitz.open(path)
        ocr_engine = RapidOCR()

        for page_index in blank_pages:
            page = doc[page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            if pix.n == 4:
                image = cv2.cvtColor(image, cv2.COLOR_RGBA2RGB)

            result, _ = ocr_engine(image)
            if not result:
                continue

            for row in result:
                text = clean_line(row[1])
                if text:
                    raw_lines.append(text)

    return raw_lines, len(reader.pages), used_ocr


def extract_docx_lines(path: Path) -> tuple[list[str], int]:
    document = Document(str(path))
    lines = [clean_line(paragraph.text) for paragraph in document.paragraphs]
    lines = [line for line in lines if line]
    return lines, len(document.paragraphs)


def extract_xlsx_lines(path: Path) -> tuple[list[str], int]:
    workbook = load_workbook(str(path), data_only=True)
    lines: list[str] = []

    for worksheet in workbook.worksheets:
        lines.append(clean_line(worksheet.title))
        for row in worksheet.iter_rows(values_only=True):
            values = [clean_line(str(cell)) for cell in row if cell is not None]
            values = [value for value in values if value]
            if values:
                lines.append(" | ".join(values))

    return [line for line in lines if line], len(workbook.worksheets)


def parse_term_definition(line: str) -> tuple[str, str] | None:
    dot_bullets = [part.strip() for part in line.split("•") if part.strip()]
    if len(dot_bullets) >= 2 and len(dot_bullets[0]) <= 120 and len(dot_bullets[1]) >= 14:
        return dot_bullets[0], dot_bullets[1]

    for separator in (" - ", " : ", ": ", " = "):
        if separator not in line:
            continue

        left, right = line.split(separator, 1)
        left = normalize_space(left)
        right = normalize_space(right)

        if len(left) < 2 or len(left) > 110:
            continue

        if len(right) < 4:
            continue

        if len(right) < 14 and separator != " = ":
            continue

        if separator == " = " and len(left.split()) > 5:
            continue

        if re.search(r"\bhttp(s)?://", left, re.IGNORECASE):
            continue

        return left, right

    return None


def shorten_term(value: str) -> str:
    value = re.sub(r"^\d+[.)\-\s]+", "", value).strip()
    value = re.sub(r"^\d+\s+", "", value).strip()
    value = re.sub(r"^\d+[.)]\s*", "", value).strip()
    value = strip_reference_noise(value)
    value = value.strip("-:;,. ")
    value = value.strip("\"'")
    value = re.sub(r"^[#*/\s]+", "", value)
    value = value.replace("*", "")
    value = normalize_space(value)

    if len(value) <= 120:
        return value

    return value[:117].rstrip() + "..."


def shorten_question(value: str) -> str:
    value = strip_reference_noise(value)
    value = value.strip("-:;,. ")
    value = normalize_space(value)

    if len(value) <= 280:
        return value

    return value[:277].rstrip() + "..."


def looks_like_code(text: str) -> bool:
    specials = sum(text.count(char) for char in "{}[];")
    if specials >= 3:
        return True

    if re.search(r"\b(public|private|protected|class|def|function|return|const|let|var)\b", text):
        if "{" in text or ";" in text or "=>" in text:
            return True

    if text.count("(") >= 2 and text.count(")") >= 2 and text.count(" ") < 12:
        return True

    return False


def is_valid_answer_label(answer: str) -> bool:
    if not answer:
        return False

    if not re.search(r"[A-Za-z]", answer):
        return False

    if len(answer) > 80:
        return False

    if not re.match(r"^[A-Za-z0-9]", answer):
        return False

    words = answer.split()
    if len(words) == 0 or len(words) > 12:
        return False

    alnum_words = [re.sub(r"[^A-Za-z0-9]", "", word) for word in words]
    if sum(1 for word in alnum_words if len(word) == 1) >= 3:
        return False

    if answer.startswith("(") and len(answer.split()) > 4:
        return False

    if re.match(r"^\d", answer):
        if re.search(r"\d.*\d", answer):
            return False

    punctuation_count = sum(1 for char in answer if not char.isalnum() and not char.isspace())
    if punctuation_count / max(1, len(answer)) > 0.2:
        return False

    noisy_words = 0
    for word in words:
        symbol_count = sum(1 for char in word if not char.isalnum() and char not in "+#.-_/")
        if symbol_count / max(1, len(word)) > 0.35:
            noisy_words += 1

    if noisy_words > 2:
        return False

    lowered = answer.lower()
    if lowered.endswith((" by", " to", " with", " and", " into", " from")):
        return False

    return True


def is_valid_question_text(question: str) -> bool:
    if len(question) < 14 or len(question) > 280:
        return False

    if re.search(r"\bhttp(s)?://", question, re.IGNORECASE):
        return False

    if question.count(";") > 2:
        return False

    alpha_count = sum(1 for char in question if char.isalpha())
    if alpha_count / max(1, len(question)) < 0.45:
        return False

    return True


def build_pairs_for_source(source_file: Path, lines: list[str]) -> list[TrainingPair]:
    pairs: list[TrainingPair] = []
    current_heading = source_file.stem.replace("_", " ")

    coalesced = coalesce_lines(lines)

    for line in coalesced:
        line = strip_reference_noise(line)
        if is_noise(line):
            continue

        if looks_like_code(line):
            continue

        if is_heading(line):
            current_heading = line
            continue

        parsed = parse_term_definition(line)
        if parsed is not None:
            term, definition = parsed
            answer = shorten_term(term)
            question = shorten_question(definition)
        else:
            answer = shorten_term(current_heading)
            question = shorten_question(line)

        if not answer or not question:
            continue

        if normalize_for_key(answer) in GENERIC_TERMS:
            continue

        if len(answer) < 2 or len(question) < 14:
            continue

        word_count = len(question.split())
        if word_count < 4:
            continue

        if question.lower().startswith("page "):
            continue

        if any(token in answer for token in ("{", "}", ";", "//")):
            continue

        if not is_valid_answer_label(answer):
            continue

        if not is_valid_question_text(question):
            continue

        pairs.append(
            TrainingPair(
                question=question,
                answer=answer,
                source_file=source_file.name,
            )
        )

    return pairs


def dedupe_pairs(pairs: list[TrainingPair]) -> list[TrainingPair]:
    def quality_rank(pair: TrainingPair) -> tuple[int, int, int, int, int]:
        answer = pair.answer
        question = pair.question

        starts_with_digit = 1 if re.match(r"^\d", answer) else 0
        punctuation_ratio = int(
            100
            * sum(1 for char in answer if not char.isalnum() and not char.isspace())
            / max(1, len(answer))
        )
        question_penalty = abs(len(question) - 120)
        short_answer_penalty = 1 if len(answer) < 4 else 0
        return (starts_with_digit, short_answer_penalty, punctuation_ratio, question_penalty, len(answer))

    by_definition: dict[str, TrainingPair] = {}

    for pair in pairs:
        definition_key = normalize_for_key(pair.question)
        pair_key = f"{definition_key}\u001f{normalize_for_key(pair.answer)}"

        if not definition_key or not normalize_for_key(pair.answer):
            continue

        existing = by_definition.get(pair_key)
        if existing is None:
            by_definition[pair_key] = pair
            continue

        if quality_rank(pair) < quality_rank(existing):
            by_definition[pair_key] = pair

    unique = list(by_definition.values())

    # Secondary dedupe: if same question appears many times with near-identical answer,
    # keep only one.
    best_by_question: dict[str, TrainingPair] = {}

    for pair in unique:
        question_key = normalize_for_key(pair.question)
        existing = best_by_question.get(question_key)
        if existing is None:
            best_by_question[question_key] = pair
            continue

        if quality_rank(pair) < quality_rank(existing):
            best_by_question[question_key] = pair

    return sorted(best_by_question.values(), key=lambda item: (item.answer.lower(), item.question.lower()))


def escape_verbatim(value: str) -> str:
    return value.replace('"', '""')


def build_seed_source_content(pairs: list[TrainingPair], generated_at: datetime, source_files: list[Path]) -> str:
    header = [
        "// <auto-generated>",
        f"// Generated by tools/extract_training_material.py on {generated_at.date().isoformat()}.",
        "// Source training files:",
    ]

    for source_file in source_files:
        header.append(f"// - {source_file}")

    header.extend(
        [
            "// </auto-generated>",
            "",
            "namespace MSSA_Jeopardy.Services;",
            "",
            "public sealed class ImportedTrainingEntry",
            "{",
            "    public string Question { get; init; } = string.Empty;",
            "    public string Answer { get; init; } = string.Empty;",
            "}",
            "",
            "public static class ImportedTrainingMaterial",
            "{",
            "    public static readonly IReadOnlyList<ImportedTrainingEntry> Entries =",
            "    [",
        ]
    )

    body: list[str] = []
    for pair in pairs:
        body.extend(
            [
                "        new ImportedTrainingEntry",
                "        {",
                f'            Question = @"{escape_verbatim(pair.question)}",',
                f'            Answer = @"{escape_verbatim(pair.answer)}"',
                "        },",
            ]
        )

    footer = [
        "    ];",
        "}",
        "",
    ]

    return "\n".join(header + body + footer)


def write_outputs(
    pairs: list[TrainingPair],
    source_stats: list[SourceStats],
    source_files: list[Path],
) -> None:
    OUTPUT_WEB_SOURCE.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_DESKTOP_SOURCE.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_REPORT.parent.mkdir(parents=True, exist_ok=True)

    now = datetime.now(timezone.utc)

    source_content = build_seed_source_content(pairs, now, source_files)
    OUTPUT_WEB_SOURCE.write_text(source_content, encoding="utf-8")
    OUTPUT_DESKTOP_SOURCE.write_text(source_content, encoding="utf-8")

    json_payload = {
        "generatedAtUtc": now.isoformat(),
        "sourceFiles": [str(path) for path in source_files],
        "pairCount": len(pairs),
        "pairs": [
            {
                "question": pair.question,
                "answer": pair.answer,
                "sourceFile": pair.source_file,
            }
            for pair in pairs
        ],
    }

    OUTPUT_JSON.write_text(json.dumps(json_payload, indent=2), encoding="utf-8")

    by_source = Counter(pair.source_file for pair in pairs)

    report_lines = [
        "# Training Extraction Report",
        "",
        f"Generated (UTC): `{now.isoformat()}`",
        "",
        f"Total unique training pairs: `{len(pairs)}`",
        "",
        "## Source Summary",
        "",
        "| Source file | Sections/pages | Extracted lines | Qualified statements | Generated pairs | OCR used |",
        "|---|---:|---:|---:|---:|:---:|",
    ]

    for stat in source_stats:
        report_lines.append(
            f"| {stat.source_file} | {stat.pages_or_sections} | {stat.extracted_lines} | {stat.qualified_statements} | {stat.generated_pairs} | {'Yes' if stat.used_ocr else 'No'} |"
        )

    report_lines.extend(
        [
            "",
            "## Output Artifacts",
            "",
            f"- Seed source (web): `{OUTPUT_WEB_SOURCE}`",
            f"- Seed source (desktop): `{OUTPUT_DESKTOP_SOURCE}`",
            f"- JSON export: `{OUTPUT_JSON}`",
            "",
            "## Placement Notes",
            "",
            "- Pairs are deduplicated across all source files using normalized question/answer keys.",
            "- The app importer maps pairs into StudyTide training tracks and lessons at seed time.",
            "",
            "## Per-source Pair Counts",
            "",
        ]
    )

    for file_name, count in sorted(by_source.items(), key=lambda item: (-item[1], item[0].lower())):
        report_lines.append(f"- `{file_name}`: `{count}` pairs")

    OUTPUT_REPORT.write_text("\n".join(report_lines) + "\n", encoding="utf-8")


def main() -> int:
    source_files = resolve_source_files()
    if not source_files:
        print("No source files found. Nothing to extract.")
        return 1

    all_pairs: list[TrainingPair] = []
    stats: list[SourceStats] = []

    for source_file in source_files:
        suffix = source_file.suffix.lower()

        if suffix == ".pdf":
            lines, page_count, used_ocr = extract_pdf_lines(source_file)
            sections = page_count
        elif suffix == ".docx":
            lines, paragraph_count = extract_docx_lines(source_file)
            sections = paragraph_count
            used_ocr = False
        elif suffix == ".xlsx":
            lines, sheet_count = extract_xlsx_lines(source_file)
            sections = sheet_count
            used_ocr = False
        else:
            continue

        coalesced = coalesce_lines(lines)
        qualified_statements = [line for line in coalesced if not is_noise(strip_reference_noise(line))]
        source_pairs = build_pairs_for_source(source_file, lines)

        all_pairs.extend(source_pairs)

        stats.append(
            SourceStats(
                source_file=source_file.name,
                pages_or_sections=sections,
                extracted_lines=len(lines),
                qualified_statements=len(qualified_statements),
                generated_pairs=len(source_pairs),
                used_ocr=used_ocr,
            )
        )

    deduped_pairs = dedupe_pairs(all_pairs)
    write_outputs(deduped_pairs, stats, source_files)

    print(f"Extracted {len(all_pairs)} raw pairs.")
    print(f"Produced {len(deduped_pairs)} unique training pairs.")
    print(f"Wrote {OUTPUT_WEB_SOURCE}")
    print(f"Wrote {OUTPUT_DESKTOP_SOURCE}")
    print(f"Wrote {OUTPUT_JSON}")
    print(f"Wrote {OUTPUT_REPORT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
